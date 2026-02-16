#!/usr/bin/env python3
"""
基于语义理解生成 Word 文档，使用 tools/create_equation.py 的方法插入公式
关键改进：
1. 先清理 LaTeX 公式，去除歧义空格，用 {} 包裹
2. 使用 MML2OMML.XSL 转换公式
"""

import sys
sys.path.insert(0, '/Users/lihengzhong/Documents/repo/devsim/tools')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

# 导入 create_equation.py 中的函数
import lxml.etree as etree
from docx.oxml import parse_xml
import latex2mathml.converter

# XSLT 转换器（全局初始化）
XSLT_PATH = '/Users/lihengzhong/Documents/repo/devsim/tools/MML2OMML.XSL'
xslt = etree.XSLT(etree.parse(XSLT_PATH))


def clean_latex(latex_str):
    """
    清理 LaTeX 公式，去除歧义空格，确保 {} 包裹
    """
    latex = latex_str.strip()
    
    # 移除行内公式标记
    if latex.startswith('$') and latex.endswith('$'):
        latex = latex[1:-1]
    
    # 1. 规范化空格：移除多余的空格，但保留必要分隔
    latex = re.sub(r'\s+', ' ', latex)
    
    # 2. 确保所有命令参数都有 {}
    # 例如：\frac kT q -> \frac{kT}{q}
    
    # 3. 处理 \frac，确保分子分母有 {}
    # 匹配 \frac{...}{...} 或 \frac a b 格式
    def fix_frac(match):
        # 如果已经是 {...}{...} 格式，保持不变
        if match.group(1).startswith('{'):
            return match.group(0)
        # 否则，尝试用空格分隔的参数包裹
        parts = match.group(1).strip().split()
        if len(parts) >= 2:
            return f"\\frac{{{parts[0]}}}{{{parts[1]}}}"
        return match.group(0)
    
    latex = re.sub(r'\\frac\s*([^\{][^\s]*(?:\s+[^\s]+)?)', fix_frac, latex)
    
    # 4. 处理下标和上标，确保有 {}
    # x_j -> x_{j}, Q_rr -> Q_{rr}, 10^-8 -> 10^{-8}
    # 匹配模式：字母_字母序列（不在{}中）
    latex = re.sub(r'([a-zA-Z])_([a-zA-Z0-9]+)(?![\{])', r'\1_{\2}', latex)
    latex = re.sub(r'([a-zA-Z0-9])\^([a-zA-Z0-9\-]+)(?![\{])', r'\1^{\2}', latex)
    
    # 5. 确保 \text{...} 中的内容不被修改
    # 但如果是 \text cm，改为 \text{cm}
    latex = re.sub(r'\\text\s+([^\{])', r'\\text{\1', latex)
    if latex.count('{') > latex.count('}'):
        latex += '}' * (latex.count('{') - latex.count('}'))
    
    # 6. 处理特殊符号
    latex = latex.replace('~', '\\sim ')
    latex = latex.replace('×', '\\times ')
    
    return latex.strip()


def add_formula_with_xslt(paragraph, latex_str, display_mode=True):
    """
    使用 XSLT 方法将 LaTeX 公式插入 Word
    """
    try:
        # 1. 清理 LaTeX
        clean_latex_str = clean_latex(latex_str)
        
        # 2. LaTeX -> MathML
        mathml = latex2mathml.converter.convert(clean_latex_str)
        
        # 3. MathML -> OMML (使用 XSLT)
        omml_tree = xslt(etree.fromstring(mathml.encode('utf-8')))
        
        # 4. 转换为字符串
        omml_str = etree.tostring(omml_tree, encoding='unicode', with_tail=False)
        if '?>' in omml_str:
            omml_str = omml_str.split('?>')[-1].strip()
        
        # 5. 添加命名空间
        m_ns = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        if 'xmlns:m=' not in omml_str:
            omml_str = omml_str.replace('<m:oMath', f'<m:oMath {m_ns}', 1)
        
        # 6. 插入段落
        if display_mode:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph._element.append(parse_xml(omml_str))
        return True
        
    except Exception as e:
        print(f"  公式转换失败: {e}")
        print(f"  LaTeX: {latex_str[:50]}...")
        # 失败时添加文本
        paragraph.add_run(f"[公式: {latex_str[:30]}...]")
        return False


def add_text_with_subscripts(para, text):
    """添加带下标的文本"""
    subscripts = [
        ('Q_{rr}', 'Q', 'rr'), ('Q_rr', 'Q', 'rr'),
        ('t_{rr}', 't', 'rr'), ('t_rr', 't', 'rr'),
        ('V_{bi}', 'V', 'bi'), ('V_bi', 'V', 'bi'),
        ('V_F', 'V', 'F'), ('V_A', 'V', 'A'),
        ('R_{on}', 'R', 'on'), ('R_on', 'R', 'on'),
        ('N_A', 'N', 'A'), ('N_D', 'N', 'D'),
        ('J_F', 'J', 'F'), ('I_F', 'I', 'F'),
        ('I_{rr}', 'I', 'rr'), ('I_rr', 'I', 'rr'),
        ('τ_n', 'τ', 'n'), ('τ_p', 'τ', 'p'),
        ('n_i', 'n', 'i'), ('x_j', 'x', 'j'),
        ('E_c', 'E', 'c'), ('E_crit', 'E', 'crit'),
        ('kT', 'k', 'T'), ('dI', 'd', 'I'), ('dV', 'd', 'V'),
    ]
    
    i = 0
    while i < len(text):
        matched = False
        for pattern, base, sub in sorted(subscripts, key=lambda x: -len(x[0])):
            if i + len(pattern) <= len(text) and text[i:i+len(pattern)] == pattern:
                run = para.add_run(base)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run = para.add_run(sub)
                run.font.subscript = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                i += len(pattern)
                matched = True
                break
        
        if not matched:
            run = para.add_run(text[i])
            run.font.name = 'Times New Roman'
            i += 1


def process_inline_text(para, text):
    """处理行内文本，识别公式、粗体和格式"""
    # 先分割粗体 **...**
    bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for bold_part in bold_parts:
        if bold_part.startswith('**') and bold_part.endswith('**') and len(bold_part) > 4:
            # 粗体文本
            bold_text = bold_part[2:-2]
            run = para.add_run(bold_text)
            run.font.bold = True
            run.font.name = 'Times New Roman'
        else:
            # 分割行内公式 $...$
            formula_parts = re.split(r'(\$[^$]+\$)', bold_part)
            
            for part in formula_parts:
                if part.startswith('$') and part.endswith('$') and len(part) > 2:
                    # 行内公式
                    latex = part[1:-1]
                    add_formula_with_xslt(para, latex, display_mode=False)
                else:
                    # 普通文本
                    if part.strip():
                        add_text_with_subscripts(para, part)


def create_docx_with_xslt():
    """基于语义生成 Word 文档，使用 XSLT 方法插入公式"""
    import re
    
    workspace_dir = '/Users/lihengzhong/Documents/repo/devsim/workspace/plan1'
    os.chdir(workspace_dir)
    
    # 读取 markdown
    with open('draft_modified_fixed.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建文档
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:]
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(title)
            run.font.name = 'SimHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            run.font.size = Pt(16)
            run.font.bold = True
            doc.add_paragraph()
        
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=1)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(14)
                run.font.bold = True
        
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=2)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(12)
                run.font.bold = True
        
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=3)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(10.5)
                run.font.bold = True
        
        # 处理行间公式
        elif line.startswith('$$') and line.endswith('$$') and len(line) > 4:
            latex = line[2:-2]
            print(f"处理公式: {latex[:50]}...")
            doc.add_paragraph()
            para = doc.add_paragraph()
            add_formula_with_xslt(para, latex, display_mode=True)
            doc.add_paragraph()
        
        elif line.startswith('$$') and not line.endswith('$$'):
            # 多行公式
            formula_lines = [line[2:]]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                formula_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                formula_lines.append(lines[i].strip()[:-2])
            latex = '\n'.join(formula_lines)
            print(f"处理多行公式: {latex[:50]}...")
            doc.add_paragraph()
            para = doc.add_paragraph()
            add_formula_with_xslt(para, latex, display_mode=True)
            doc.add_paragraph()
        
        # 处理图片
        elif line.startswith('!['):
            import re
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match and os.path.exists(match.group(2)):
                doc.add_paragraph()
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(match.group(2), width=Inches(5.5))
                
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(match.group(1))
                run.font.size = Pt(9)
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                doc.add_paragraph()
        
        # 处理表格标题
        elif line.startswith('**表') and line.endswith('**'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line[2:-2])
            run.font.bold = True
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # 处理表格
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 3:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                rows = [[c.strip() for c in row.split('|')[1:-1]] for row in table_lines[2:] if row.strip()]
                
                if headers and rows:
                    table = doc.add_table(rows=1+len(rows), cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    for j, h in enumerate(headers):
                        table.rows[0].cells[j].text = h
                        for p in table.rows[0].cells[j].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                r.font.size = Pt(9)
                    
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            if ci < len(headers):
                                table.rows[ri+1].cells[ci].text = cell
                                for p in table.rows[ri+1].cells[ci].paragraphs:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for r in p.runs:
                                        r.font.size = Pt(9)
                    
                    doc.add_paragraph()
            continue
        
        # 处理列表
        elif re.match(r'^[\-\*]\s+', line):
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            text = re.sub(r'^[\-\*]\s+', '', line)
            process_inline_text(para, text)
        
        # 处理参考文献
        elif re.match(r'^\[\d+\]', line):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            process_inline_text(para, line)
        
        # 普通段落
        else:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            process_inline_text(para, line)
        
        i += 1
    
    # 保存
    output_path = '功率二极管反向恢复特性研究_语义生成.docx'
    doc.save(output_path)
    print(f"\n✅ Word文档已生成: {output_path}")
    print(f"📄 文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
    
    return output_path


if __name__ == '__main__':
    print("="*70)
    print("使用 XSLT 方法生成 Word 文档")
    print("="*70)
    print()
    
    try:
        create_docx_with_xslt()
        print()
        print("="*70)
        print("生成完成！使用 MML2OMML.XSL 转换公式。")
        print("="*70)
    except Exception as e:
        print(f"\n❌ 生成失败: {e}")
        import traceback
        traceback.print_exc()
