#!/usr/bin/env python3
"""
对比测试：MathML 包装 vs OMML 转换 两种公式生成方式
生成一个 Word 文档，同时包含两种方式生成的公式，便于对比
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import latex2mathml.converter
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml


def add_math_method1_mathml_wrap(paragraph, latex_str):
    """
    方式1：MathML 包装法
    直接将 LaTeX 转为 MathML，包装在 <m:oMath> 中
    """
    try:
        # LaTeX -> MathML
        mathml = latex2mathml.converter.convert(latex_str)
        
        # 包装在 OMML 容器中
        namespace = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        omml_xml = f'<m:oMath {namespace}>{mathml}</m:oMath>'
        
        # 插入
        element = parse_xml(omml_xml)
        paragraph._element.append(element)
        return True
    except Exception as e:
        print(f"  方式1失败: {e}")
        return False


def add_math_method2_omml_convert(paragraph, latex_str):
    """
    方式2：OMML 转换法
    LaTeX -> MathML -> OMML，生成原生 Word 公式
    """
    try:
        # LaTeX -> MathML -> OMML
        mathml = latex_to_mathml(latex_str)
        omml = mathml_to_omml(mathml)
        
        # 添加命名空间
        omml_with_ns = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</m:oMath>'
        
        # 插入
        element = parse_xml(omml_with_ns)
        paragraph._element.append(element)
        return True
    except Exception as e:
        print(f"  方式2失败: {e}")
        return False


def add_text_formula(paragraph, latex_str):
    """
    方式3：文本降级（备用方案）
    显示为带下标的文本
    """
    import re
    from docx.shared import Pt
    
    # 清理 LaTeX
    text = latex_str
    text = text.replace('\\text{', '').replace('}', '')
    text = text.replace('\\times', '×').replace('\\cdot', '·')
    text = text.replace('\\sim', '~').replace('\\approx', '≈')
    text = text.replace('\\leq', '≤').replace('\\geq', '≥')
    text = text.replace('\\', '')
    
    # 处理下标
    i = 0
    while i < len(text):
        sub_match = re.match(r'([a-zA-Zτ])_\{?([a-zA-Z0-9]+)\}?', text[i:])
        if sub_match:
            base = sub_match.group(1)
            sub = sub_match.group(2)
            run = paragraph.add_run(base)
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run = paragraph.add_run(sub)
            run.font.subscript = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            i += sub_match.end()
        elif text[i:i+2] == '10' and i+2 < len(text) and text[i+2] == '^':
            j = i + 3
            while j < len(text) and (text[j].isdigit() or text[j] in '{-}'):
                j += 1
            exp = text[i+3:j].replace('{', '').replace('}', '')
            run = paragraph.add_run('10')
            run.font.name = 'Times New Roman'
            run = paragraph.add_run(exp)
            run.font.superscript = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            i = j
        else:
            run = paragraph.add_run(text[i])
            run.font.name = 'Times New Roman'
            i += 1


def create_comparison_doc():
    """生成对比测试文档"""
    
    doc = Document()
    
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('公式生成方式对比测试', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(16)
        run.font.bold = True
    
    doc.add_paragraph()
    
    # 测试公式列表
    test_formulas = [
        ("简单下标", r"N_A = 10^{16} \text{ cm}^{-3}"),
        ("范围表示", r"\tau = 10^{-8} \sim 10^{-4} \text{ s}"),
        ("复杂分数", r"V_{bi} = \frac{kT}{q} \ln\left(\frac{N_A N_D}{n_i^2}\right)"),
        ("根号", r"R_{on} \propto \frac{1}{\sqrt{N_A}}"),
        ("近似关系", r"Q_{rr} \approx 3.0 \cdot \tau \cdot J_F"),
        ("多下标", r"I_{rr} = f(\tau_n, \tau_p, N_A, N_D)"),
    ]
    
    for name, latex in test_formulas:
        # 公式名称
        heading = doc.add_heading(name, level=2)
        for run in heading.runs:
            run.font.name = 'SimHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            run.font.size = Pt(12)
            run.font.bold = True
        
        # LaTeX 源码
        p_src = doc.add_paragraph()
        p_src.add_run("LaTeX: ").font.bold = True
        p_src.add_run(latex).font.name = 'Courier New'
        
        # 方式1：MathML 包装
        doc.add_paragraph("【方式1】MathML 包装法：")
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        success1 = add_math_method1_mathml_wrap(p1, latex)
        if not success1:
            p1.add_run("[生成失败，降级为文本]")
            add_text_formula(p1, latex)
        
        # 方式2：OMML 转换
        doc.add_paragraph("【方式2】OMML 转换法：")
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        success2 = add_math_method2_omml_convert(p2, latex)
        if not success2:
            p2.add_run("[生成失败，降级为文本]")
            add_text_formula(p2, latex)
        
        # 方式3：文本显示（对比）
        doc.add_paragraph("【方式3】文本格式（对比）：")
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_formula(p3, latex)
        
        doc.add_paragraph()  # 空行分隔
    
    # 保存
    output_path = '公式生成方式对比测试.docx'
    doc.save(output_path)
    print(f"✅ 对比测试文档已生成: {output_path}")
    
    # 统计信息
    import os
    import zipfile
    import re
    
    file_size = os.path.getsize(output_path) / 1024
    print(f"📄 文件大小: {file_size:.1f} KB")
    
    # 分析生成的 XML
    with zipfile.ZipFile(output_path, 'r') as z:
        xml = z.read('word/document.xml').decode('utf-8')
        
        # 统计公式数量
        mathml_count = len(re.findall(r'<math xmlns=', xml))
        omml_count = len(re.findall(r'<m:r>', xml))
        
        print(f"\n📊 公式统计：")
        print(f"  - MathML 包装公式: {mathml_count}")
        print(f"  - OMML 元素数量: {omml_count}")
    
    print(f"\n📝 测试说明：")
    print(f"  请在 Word/WPS 中打开此文件，对比三种方式的显示效果：")
    print(f"  1. MathML 包装法 - 通用性强")
    print(f"  2. OMML 转换法 - Word 原生格式")
    print(f"  3. 文本格式 - 备用方案")
    
    return output_path


if __name__ == '__main__':
    create_comparison_doc()
    print("\n" + "="*70)
    print("对比测试完成！请在 Word 中检查显示效果。")
    print("="*70)
