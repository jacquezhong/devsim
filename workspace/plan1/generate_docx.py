#!/usr/bin/env python3
"""
从 draft_modified.md 生成 功率二极管反向恢复特性研究.docx

公式处理策略：
1. 简单下标/上标（如 N_A, τ_n, 10^14）→ Word 公式对象（可编辑）
2. 复杂公式（含 \frac, \ln, 嵌套括号等）→ 清晰的文本格式，用户可手动在Word中按 Alt+= 转换

这样确保：
- 简单符号正确显示
- 复杂公式不显示为乱码
- 用户可在Word中手动编辑复杂公式
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os


def is_simple_formula(latex):
    """
    判断是否为简单公式（只包含下标、上标、基本运算符）
    简单公式可以安全转换为 Word OMML
    """
    # 清理
    latex = latex.strip()
    
    # 如果有这些复杂命令，认为是复杂公式
    complex_commands = [
        '\\frac', '\\sqrt', '\\ln', '\\log', '\\sin', '\\cos', '\\tan',
        '\\int', '\\sum', '\\prod', '\\lim', '\\begin', '\\end',
        '\\left', '\\right', '\\overline', '\\underline',
        '\\hat', '\\bar', '\\vec', '\\dot', '\\ddot',
    ]
    
    for cmd in complex_commands:
        if cmd in latex:
            return False
    
    # 检查是否有复杂的嵌套括号
    # 简单公式可以包含括号，但不能有复杂的嵌套结构
    depth = 0
    max_depth = 0
    for char in latex:
        if char == '(' or char == '[' or char == '{':
            depth += 1
            max_depth = max(max_depth, depth)
        elif char == ')' or char == ']' or char == '}':
            depth -= 1
    
    # 如果最大嵌套深度超过1，认为是复杂公式
    if max_depth > 1:
        return False
    
    return True


def add_simple_formula_omml(para, latex):
    """
    将简单公式转换为 Word OMML
    只处理：变量、数字、下标、上标、基本运算符
    """
    # 预定义替换
    latex = latex.replace('\\times', '×').replace('\\cdot', '·')
    latex = latex.replace('\\approx', '≈').replace('\\propto', '∝')
    latex = latex.replace('\\leq', '≤').replace('\\geq', '≥')
    latex = latex.replace('\\', '')  # 移除剩余反斜杠
    latex = latex.replace('{', '').replace('}', '')
    latex = latex.replace('text', '')
    
    # 转义 XML
    latex = latex.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # 构建 OMML
    parts = []
    i = 0
    n = len(latex)
    
    while i < n:
        # 跳过空格
        if latex[i].isspace():
            i += 1
            continue
        
        # 检查下标模式：X_Y
        sub_match = re.match(r'([a-zA-Zταβγδεθλμπσφωρηκ])_([a-zA-Z0-9]+)', latex[i:])
        if sub_match:
            base = sub_match.group(1)
            sub = sub_match.group(2)
            parts.append(f'<m:sSub><m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{base}</m:t></m:r></m:e><m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub></m:sSub>')
            i += sub_match.end()
            continue
        
        # 检查上标模式：10^n 或 X^n
        sup_match = re.match(r'(\d+(?:\.\d+)?)?×?10\^(-?\d+)', latex[i:])
        if sup_match:
            coeff = sup_match.group(1) or ''
            exp = sup_match.group(2)
            if coeff:
                parts.append(f'<m:r><m:t>{coeff}×10</m:t></m:r>')
            else:
                parts.append(f'<m:r><m:t>10</m:t></m:r>')
            parts.append(f'<m:sSup><m:e/><m:sup><m:r><m:t>{exp}</m:t></m:r></m:sup></m:sSup>')
            i += sup_match.end()
            continue
        
        # 希腊字母（斜体）
        if latex[i] in 'ταβγδεθλμπσφωρηκ':
            parts.append(f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{latex[i]}</m:t></m:r>')
            i += 1
            continue
        
        # 英文字母变量（斜体）
        if latex[i].isalpha():
            parts.append(f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{latex[i]}</m:t></m:r>')
            i += 1
            continue
        
        # 数字（正体）
        if latex[i].isdigit() or latex[i] == '.':
            j = i
            while j < n and (latex[j].isdigit() or latex[j] == '.'):
                j += 1
            parts.append(f'<m:r><m:t>{latex[i:j]}</m:t></m:r>')
            i = j
            continue
        
        # 运算符和符号
        if latex[i] in '+-=×·≈∝≤≥(),/;':
            parts.append(f'<m:r><m:t>{latex[i]}</m:t></m:r>')
            i += 1
            continue
        
        # 其他字符
        parts.append(f'<m:r><m:t>{latex[i]}</m:t></m:r>')
        i += 1
    
    if parts:
        try:
            from docx.oxml import parse_xml
            omml_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{ "".join(parts) }</m:oMath>'
            element = parse_xml(omml_xml)
            para._p.append(element)
            return True
        except Exception as e:
            print(f"  OMML构建失败: {e}")
    
    return False


def add_formula_as_formatted_text(para, latex):
    """
    将公式显示为格式化的文本（清晰的下标/上标）
    用户可以在Word中手动选中并按 Alt+= 转换为公式
    """
    # 清理但保留可读性
    latex = latex.replace('\\times', '×').replace('\\cdot', '·')
    latex = latex.replace('\\approx', '≈').replace('\\propto', '∝')
    latex = latex.replace('\\leq', '≤').replace('\\geq', '≥')
    latex = latex.replace('\\left', '').replace('\\right', '')
    latex = latex.replace('\\ln', 'ln').replace('\\log', 'log')
    latex = latex.replace('\\sin', 'sin').replace('\\cos', 'cos').replace('\\tan', 'tan')
    latex = latex.replace('\\frac', '').replace('\\sqrt', '√')
    latex = latex.replace('\\begin', '').replace('\\end', '').replace('cases', '')
    latex = latex.replace('\\', '')
    latex = latex.replace('{', '').replace('}', '')
    latex = latex.replace('text', '')
    latex = latex.replace('\\\\', '; ')
    
    # 定义符号模式 (模式, 基础, 下标)
    subscript_patterns = [
        (r'Q_{rr}', 'Q', 'rr'), (r'Q_rr', 'Q', 'rr'),
        (r't_{rr}', 't', 'rr'), (r't_rr', 't', 'rr'),
        (r'V_{bi}', 'V', 'bi'), (r'V_bi', 'V', 'bi'),
        (r'V_F', 'V', 'F'), (r'V_A', 'V', 'A'),
        (r'R_{on}', 'R', 'on'), (r'R_on', 'R', 'on'),
        (r'N_A', 'N', 'A'), (r'N_D', 'N', 'D'),
        (r'J_F', 'J', 'F'), (r'I_F', 'I', 'F'),
        (r'I_{rr}', 'I', 'rr'), (r'I_rr', 'I', 'rr'),
        (r'τ_n', 'τ', 'n'), (r'τ_p', 'τ', 'p'),
        (r'n_i', 'n', 'i'), (r'x_j', 'x', 'j'),
        (r'E_c', 'E', 'c'), (r'E_crit', 'E', 'crit'),
        (r'kT', 'k', 'T'), (r'dI', 'd', 'I'), (r'dV', 'd', 'V'),
    ]
    
    i = 0
    text = latex
    while i < len(text):
        matched = False
        
        # 检查下标模式
        for pattern, base, sub in sorted(subscript_patterns, key=lambda x: -len(x[0])):
            p = pattern.replace('\\', '').replace('{', '').replace('}', '')
            if i + len(p) <= len(text) and text[i:i+len(p)] == p:
                # 基础符号（斜体）
                run = para.add_run(base)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                # 下标
                run = para.add_run(sub)
                run.font.subscript = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                i += len(p)
                matched = True
                break
        
        if not matched:
            # 科学计数法上标
            sup_match = re.match(r'10\^(-?\d+)', text[i:])
            if sup_match:
                run = para.add_run('10')
                run.font.name = 'Times New Roman'
                run = para.add_run(sup_match.group(1))
                run.font.superscript = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                i += sup_match.end()
                matched = True
            
            if not matched:
                # 希腊字母（斜体）
                if text[i] in 'ταβγδεθλμπσφωρηκ':
                    run = para.add_run(text[i])
                    run.font.italic = True
                    run.font.name = 'Times New Roman'
                    i += 1
                    matched = True
                
                if not matched:
                    # 英文字母（斜体）
                    if text[i].isalpha():
                        run = para.add_run(text[i])
                        run.font.italic = True
                        run.font.name = 'Times New Roman'
                        i += 1
                        matched = True
                    
                    if not matched:
                        # 其他字符
                        run = para.add_run(text[i])
                        run.font.name = 'Times New Roman'
                        i += 1


def add_formula(para, latex_str, display_mode=True):
    """
    添加公式到段落
    根据复杂度决定使用 OMML 还是格式化文本
    """
    latex = latex_str.strip()
    if latex.startswith('$') and latex.endswith('$'):
        latex = latex[1:-1]
    
    if display_mode:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 判断是否为简单公式
    if is_simple_formula(latex):
        # 尝试转换为 OMML
        if add_simple_formula_omml(para, latex):
            return True
    
    # 使用格式化文本
    add_formula_as_formatted_text(para, latex)
    return False


def process_inline_formatting(para, text):
    """处理行内格式（粗体、斜体、公式）"""
    # 分割粗体、斜体、公式
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\$[^$]+\$)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # 粗体
            run = para.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Times New Roman'
        elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
            # 斜体
            run = para.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Times New Roman'
        elif part.startswith('$') and part.endswith('$') and len(part) > 2:
            # 行内公式
            add_formula(para, part[1:-1], display_mode=False)
        else:
            # 普通文本（处理下标）
            if part.strip():
                add_text_with_subscripts(para, part)


def add_text_with_subscripts(para, text):
    """处理文本中的下标"""
    patterns = [
        (r'Q_{rr}', 'Q', 'rr'), (r'Q_rr', 'Q', 'rr'),
        (r't_{rr}', 't', 'rr'), (r't_rr', 't', 'rr'),
        (r'V_{bi}', 'V', 'bi'), (r'V_bi', 'V', 'bi'),
        (r'V_F', 'V', 'F'), (r'V_A', 'V', 'A'),
        (r'R_{on}', 'R', 'on'), (r'R_on', 'R', 'on'),
        (r'N_A', 'N', 'A'), (r'N_D', 'N', 'D'),
        (r'J_F', 'J', 'F'), (r'I_F', 'I', 'F'),
        (r'I_{rr}', 'I', 'rr'), (r'I_rr', 'I', 'rr'),
        (r'τ_n', 'τ', 'n'), (r'τ_p', 'τ', 'p'),
        (r'n_i', 'n', 'i'), (r'x_j', 'x', 'j'),
        (r'E_c', 'E', 'c'), (r'E_crit', 'E', 'crit'),
        (r'kT', 'k', 'T'), (r'dI', 'd', 'I'), (r'dV', 'd', 'V'),
    ]
    
    i = 0
    while i < len(text):
        matched = False
        
        for pattern, base, sub in sorted(patterns, key=lambda x: -len(x[0])):
            p = pattern.replace('\\', '').replace('{', '').replace('}', '')
            if i + len(p) <= len(text) and text[i:i+len(p)] == p:
                run = para.add_run(base)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run = para.add_run(sub)
                run.font.subscript = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                i += len(p)
                matched = True
                break
        
        if not matched:
            if i < len(text) - 1 and text[i] == '^' and (text[i+1].isdigit() or text[i+1] in '+-'):
                j = i + 2 if text[i+1] in '+-' else i + 1
                while j < len(text) and text[j].isdigit():
                    j += 1
                if j > i + 1:
                    run = para.add_run(text[i+1:j])
                    run.font.superscript = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    i = j
                    matched = True
            
            if not matched:
                run = para.add_run(text[i])
                run.font.name = 'Times New Roman'
                i += 1


def create_docx():
    """主函数"""
    workspace_dir = '/Users/lihengzhong/Documents/repo/devsim/workspace/plan1'
    os.chdir(workspace_dir)
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    # 使用 draft_modified.md
    with open('draft_modified.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            i += 1
            continue
        
        # 标题处理
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(16)
                run.font.bold = True
            doc.add_paragraph()
        
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=1)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(14)
                run.font.bold = True
        
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=2)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(12)
                run.font.bold = True
        
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=3)
            for run in heading.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(10.5)
                run.font.bold = True
        
        # 图片
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match and os.path.exists(match.group(2)):
                doc.add_paragraph()
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(match.group(2), width=Inches(5.5))
                
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(match.group(1))
                run.font.size = Pt(9)
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                doc.add_paragraph()
        
        # 行间公式
        elif line.startswith('$$') and line.endswith('$$') and len(line) > 4:
            doc.add_paragraph()
            para = doc.add_paragraph()
            add_formula(para, line[2:-2], display_mode=True)
            doc.add_paragraph()
        
        # 多行公式
        elif line.startswith('$$') and not line.endswith('$$'):
            formula_lines = [line[2:]]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                formula_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                formula_lines.append(lines[i].strip()[:-2])
            doc.add_paragraph()
            para = doc.add_paragraph()
            add_formula(para, '\n'.join(formula_lines), display_mode=True)
            doc.add_paragraph()
        
        # 表格标题
        elif line.startswith('**表') and line.endswith('**'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line[2:-2])
            run.font.bold = True
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # 表格
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 3:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                rows = [[c.strip() for c in row.split('|')[1:-1]] for row in table_lines[2:] if row.strip()]
                
                if headers and rows:
                    table = doc.add_table(rows=1+len(rows), cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    for j, h in enumerate(headers):
                        table.rows[0].cells[j].text = h
                        for p in table.rows[0].cells[j].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                r.font.size = Pt(9)
                    
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            if ci < len(headers):
                                table.rows[ri+1].cells[ci].text = cell
                                for p in table.rows[ri+1].cells[ci].paragraphs:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for r in p.runs:
                                        r.font.size = Pt(9)
                    
                    doc.add_paragraph()
            continue
        
        # 列表
        elif re.match(r'^[\-\*]\s+\*\*', line):
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            process_inline_formatting(para, re.sub(r'^[\-\*]\s+', '', line))
        
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            process_inline_formatting(para, line[2:])
        
        # 参考文献
        elif re.match(r'^\[\d+\]', line):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            process_inline_formatting(para, line)
        
        # 普通段落
        else:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            process_inline_formatting(para, line)
        
        i += 1
    
    output_path = '功率二极管反向恢复特性研究.docx'
    doc.save(output_path)
    print(f"✅ Word文档已生成: {output_path}")
    print(f"📄 文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
    
    return output_path


if __name__ == '__main__':
    create_docx()
    print("\n" + "="*70)
    print("生成完成！复杂公式显示为格式化文本，可在Word中手动编辑。")
    print("提示：选中公式文本，按 Alt+= 可转换为Word公式对象。")
    print("="*70)
