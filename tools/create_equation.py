import lxml.etree as etree
from docx import Document
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
import latex2mathml.converter
import re
import os

# XSLT 文件路径（相对于本文件）
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
XSLT_PATH = os.path.join(CURRENT_DIR, "MML2OMML.XSL")

# 全局 XSLT 转换器（避免重复加载）
_xslt = None

def _get_xslt():
    """获取或初始化 XSLT 转换器"""
    global _xslt
    if _xslt is None:
        _xslt = etree.XSLT(etree.parse(XSLT_PATH))
    return _xslt

def clean_latex(latex_str):
    """
    清理 LaTeX 公式，处理多字符下标/上标
    
    Args:
        latex_str: LaTeX 字符串
    
    Returns:
        清理后的 LaTeX 字符串
    """
    latex = latex_str.strip()
    
    # 移除行内公式标记
    if latex.startswith('$') and latex.endswith('$'):
        latex = latex[1:-1]
    
    # 处理多字母下标：Q_rr -> Q_{rr}
    latex = re.sub(r'([a-zA-Z])_([a-zA-Z0-9]+)(?![\{])', r'\1_{\2}', latex)
    
    # 处理上标：10^-8 -> 10^{-8}
    latex = re.sub(r'([a-zA-Z0-9])\^([a-zA-Z0-9\-]+)(?![\{])', r'\1^{\2}', latex)
    
    # 规范化特殊符号
    latex = latex.replace('~', '\\sim ')
    latex = latex.replace('×', '\\times ')
    
    return latex

def add_formula(paragraph, latex_str, display_mode=True):
    """
    使用 XSLT 方法将 LaTeX 公式插入 Word 段落
    
    Args:
        paragraph: docx Paragraph 对象
        latex_str: LaTeX 公式字符串
        display_mode: 是否为显示模式（居中）
    
    Returns:
        bool: 是否成功
    """
    try:
        # 1. 清理 LaTeX
        clean_latex_str = clean_latex(latex_str)
        
        # 2. LaTeX -> MathML
        mathml = latex2mathml.converter.convert(clean_latex_str)
        
        # 3. MathML -> OMML (使用 XSLT)
        xslt = _get_xslt()
        omml_tree = xslt(etree.fromstring(mathml.encode('utf-8')))
        
        # 4. 转换为字符串并添加命名空间
        omml_str = etree.tostring(omml_tree, encoding='unicode', with_tail=False)
        if '?>' in omml_str:
            omml_str = omml_str.split('?>')[-1].strip()
        
        m_ns = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        if 'xmlns:m=' not in omml_str:
            omml_str = omml_str.replace('<m:oMath', f'<m:oMath {m_ns}', 1)
        
        # 5. 插入文档
        if display_mode:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph._element.append(parse_xml(omml_str))
        return True
        
    except Exception as e:
        print(f"公式转换失败: {e}")
        paragraph.add_run(f"[公式: {latex_str[:30]}...]")
        return False

def add_simple_omml(doc, latex_str):
    """
    简化版：直接向文档添加公式（新建段落）
    
    Args:
        doc: Document 对象
        latex_str: LaTeX 公式字符串
    """
    p = doc.add_paragraph()
    add_formula(p, latex_str, display_mode=True)

if __name__ == '__main__':
    # --- 运行测试 ---
    doc = Document()
    
    # 希腊字母测试
    latex_eq = r"\alpha + \beta = \sqrt{\gamma^2 + \delta^2}\cdot \sum_{i=1}^{n}{x_i}"
    add_simple_omml(doc, latex_eq)
    
    # 半导体常用公式测试
    p = doc.add_paragraph()
    add_formula(p, r'Q_{rr} = \tau \cdot J_F')
    
    try:
        doc.save("equation_test.docx")
        print("生成成功！请检查 equation_test.docx")
    except Exception as e:
        print(f"失败原因: {e}")
